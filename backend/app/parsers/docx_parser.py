"""DocxParser - 基于 python-docx，纯 Python 兼容 LoongArch."""

from __future__ import annotations

import io

from app.parsers.base import ParsedDocument, ParserError, TitleNode


class DocxParser:
    file_type = "docx"

    async def parse(self, content: bytes) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError as e:
            raise ParserError("python-docx 未安装", field="dependency") from e

        try:
            doc = Document(io.BytesIO(content))
        except Exception as e:  # noqa: BLE001
            raise ParserError(f"docx 解析失败: {e}", field="file") from e

        paragraphs: list[str] = []
        title_tree: list[TitleNode] = []
        title_stack: list[TitleNode] = []
        raw_text_parts: list[str] = []

        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            raw_text_parts.append(text)
            style = (para.style.name if para.style else "") or ""
            level = _heading_level(style)
            if level > 0:
                node = TitleNode(level=level, text=text, children=[])
                # 维护栈
                while title_stack and title_stack[-1].level >= level:
                    title_stack.pop()
                if not title_stack:
                    title_tree.append(node)
                else:
                    title_stack[-1].children.append(node)
                title_stack.append(node)
            else:
                paragraphs.append(text)

        # 表格
        tables: list[list[list[str]]] = []
        for tbl in doc.tables:
            rows: list[list[str]] = []
            for row in tbl.rows:
                rows.append([(c.text or "").strip() for c in row.cells])
            tables.append(rows)
            for row in rows:
                raw_text_parts.append(" | ".join(row))

        return ParsedDocument(
            raw_text="\n".join(raw_text_parts),
            paragraphs=paragraphs,
            title_tree=title_tree,
            tables=tables,
            page_count=0,  # docx 无显式页数
        )


def _heading_level(style_name: str) -> int:
    """从样式名提取标题层级；非标题返回 0."""
    if style_name.startswith("Heading "):
        try:
            return int(style_name[len("Heading ") :])
        except ValueError:
            return 0
    if style_name.startswith("标题 "):
        try:
            return int(style_name[len("标题 ") :])
        except ValueError:
            return 0
    return 0
