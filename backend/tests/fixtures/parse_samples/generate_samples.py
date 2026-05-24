"""生成各种格式的测试样本文件，用于验证解析链路.

运行方式：
    cd backend
    python tests/fixtures/parse_samples/generate_samples.py

生成的文件：
1. sample_report.docx - 完整的实训报告（Word）
2. sample_database_design.pdf - 数据库设计文档（PDF）
3. sample_data_analysis.xlsx - 数据分析表格（Excel）
4. sample_flask_project.zip - Flask Web 项目源代码
5. sample_screenshot.png - 程序运行截图（含中文文字）
6. sample_algorithm.docx - 算法实验报告
7. sample_network_config.pdf - 网络配置文档
8. sample_ml_notebook.xlsx - 机器学习实验数据
9. sample_vue_project.zip - Vue.js 前端项目
10. sample_api_test.docx - API 测试报告
"""

from __future__ import annotations

import io
import os
import zipfile
from pathlib import Path

OUTPUT_DIR = Path(__file__).parent / "output"
OUTPUT_DIR.mkdir(exist_ok=True)


def generate_docx_report() -> None:
    """生成实训报告 Word 文档."""
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    # 标题
    doc.add_heading("Python Web 开发实训报告", level=0)
    doc.add_paragraph("学生姓名：李明  学号：2024001  班级：软件2401")
    doc.add_paragraph("指导教师：王伟  提交日期：2024-12-15")
    doc.add_paragraph("")

    # 第一章
    doc.add_heading("第一章 项目概述", level=1)
    doc.add_heading("1.1 项目背景", level=2)
    doc.add_paragraph(
        "本项目旨在开发一个基于 Flask 框架的在线图书管理系统。"
        "系统支持图书的增删改查、用户注册登录、借阅管理等核心功能。"
        "采用 MVC 架构模式，前后端分离设计，数据库使用 MySQL 5.7。"
    )
    doc.add_heading("1.2 技术选型", level=2)
    doc.add_paragraph("后端框架：Flask 2.3.2")
    doc.add_paragraph("数据库：MySQL 5.7 + SQLAlchemy ORM")
    doc.add_paragraph("前端：Bootstrap 5 + jQuery 3.6")
    doc.add_paragraph("部署：Gunicorn + Nginx")

    # 第二章
    doc.add_heading("第二章 需求分析", level=1)
    doc.add_heading("2.1 功能需求", level=2)
    doc.add_paragraph("1. 用户管理：注册、登录、个人信息修改、密码重置")
    doc.add_paragraph("2. 图书管理：图书录入、编辑、删除、分类管理")
    doc.add_paragraph("3. 借阅管理：借书、还书、续借、逾期提醒")
    doc.add_paragraph("4. 搜索功能：按书名、作者、ISBN 搜索")
    doc.add_paragraph("5. 统计报表：借阅排行、热门图书、用户活跃度")

    doc.add_heading("2.2 非功能需求", level=2)
    doc.add_paragraph("- 系统响应时间 < 2 秒")
    doc.add_paragraph("- 支持 50 用户并发访问")
    doc.add_paragraph("- 数据库备份每日自动执行")
    doc.add_paragraph("- 密码使用 bcrypt 加密存储")

    # 第三章
    doc.add_heading("第三章 系统设计", level=1)
    doc.add_heading("3.1 数据库设计", level=2)
    doc.add_paragraph(
        "系统共设计 5 张核心数据表：users、books、categories、"
        "borrow_records、reviews。各表之间通过外键关联，"
        "保证数据一致性和完整性。"
    )

    # 添加表格
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers = ["字段名", "类型", "约束", "说明"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["id", "INT", "PRIMARY KEY", "用户ID"],
        ["username", "VARCHAR(50)", "UNIQUE NOT NULL", "用户名"],
        ["password_hash", "VARCHAR(128)", "NOT NULL", "密码哈希"],
        ["email", "VARCHAR(100)", "UNIQUE", "邮箱"],
        ["created_at", "DATETIME", "DEFAULT NOW()", "创建时间"],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = val

    doc.add_heading("3.2 API 接口设计", level=2)
    doc.add_paragraph("RESTful API 设计遵循以下规范：")
    doc.add_paragraph("- GET /api/books - 获取图书列表（支持分页）")
    doc.add_paragraph("- POST /api/books - 新增图书")
    doc.add_paragraph("- PUT /api/books/{id} - 更新图书信息")
    doc.add_paragraph("- DELETE /api/books/{id} - 删除图书")
    doc.add_paragraph("- POST /api/auth/login - 用户登录")
    doc.add_paragraph("- POST /api/borrow - 借阅图书")

    # 第四章
    doc.add_heading("第四章 系统实现", level=1)
    doc.add_heading("4.1 核心代码实现", level=2)
    doc.add_paragraph(
        "以下为图书借阅核心逻辑的实现代码片段："
    )
    doc.add_paragraph(
        "```python\n"
        "@app.route('/api/borrow', methods=['POST'])\n"
        "@login_required\n"
        "def borrow_book():\n"
        "    book_id = request.json.get('book_id')\n"
        "    book = Book.query.get_or_404(book_id)\n"
        "    if book.available_copies <= 0:\n"
        "        return jsonify({'error': '库存不足'}), 400\n"
        "    record = BorrowRecord(\n"
        "        user_id=current_user.id,\n"
        "        book_id=book_id,\n"
        "        borrow_date=datetime.now(),\n"
        "        due_date=datetime.now() + timedelta(days=30)\n"
        "    )\n"
        "    book.available_copies -= 1\n"
        "    db.session.add(record)\n"
        "    db.session.commit()\n"
        "    return jsonify({'message': '借阅成功'}), 201\n"
        "```"
    )

    # 第五章
    doc.add_heading("第五章 测试与部署", level=1)
    doc.add_heading("5.1 单元测试", level=2)
    doc.add_paragraph("使用 pytest 编写了 45 个单元测试用例，覆盖率达到 82%。")
    doc.add_paragraph("测试覆盖了用户认证、图书 CRUD、借阅逻辑等核心模块。")
    doc.add_heading("5.2 部署方案", level=2)
    doc.add_paragraph(
        "生产环境使用 Docker Compose 部署，包含 Web 服务、"
        "MySQL 数据库、Redis 缓存三个容器。"
    )

    # 总结
    doc.add_heading("总结与展望", level=1)
    doc.add_paragraph(
        "本次实训完成了图书管理系统的全部核心功能开发，"
        "掌握了 Flask Web 开发、数据库设计、RESTful API 设计、"
        "单元测试和 Docker 部署等技能。后续计划增加推荐算法和移动端适配。"
    )

    doc.save(str(OUTPUT_DIR / "sample_report.docx"))
    print("✓ sample_report.docx")


def generate_docx_algorithm() -> None:
    """生成算法实验报告 Word 文档."""
    from docx import Document

    doc = Document()
    doc.add_heading("数据结构与算法实验报告", level=0)
    doc.add_paragraph("实验名称：排序算法性能对比分析")
    doc.add_paragraph("学生：张文卓  学号：2024002")
    doc.add_paragraph("")

    doc.add_heading("一、实验目的", level=1)
    doc.add_paragraph(
        "1. 理解并实现冒泡排序、快速排序、归并排序、堆排序四种排序算法\n"
        "2. 通过实验对比不同规模数据下各算法的时间复杂度\n"
        "3. 分析各算法的空间复杂度和稳定性"
    )

    doc.add_heading("二、实验环境", level=1)
    doc.add_paragraph("- 操作系统：Ubuntu 22.04 LTS")
    doc.add_paragraph("- 编程语言：Python 3.11")
    doc.add_paragraph("- 测试数据规模：1000, 5000, 10000, 50000, 100000")

    doc.add_heading("三、算法实现", level=1)
    doc.add_heading("3.1 快速排序", level=2)
    doc.add_paragraph(
        "快速排序采用分治策略，选取基准元素将数组分为两部分，"
        "递归排序子数组。平均时间复杂度 O(n log n)，最坏 O(n²)。"
    )
    doc.add_paragraph(
        "```python\n"
        "def quicksort(arr: list[int]) -> list[int]:\n"
        "    if len(arr) <= 1:\n"
        "        return arr\n"
        "    pivot = arr[len(arr) // 2]\n"
        "    left = [x for x in arr if x < pivot]\n"
        "    middle = [x for x in arr if x == pivot]\n"
        "    right = [x for x in arr if x > pivot]\n"
        "    return quicksort(left) + middle + quicksort(right)\n"
        "```"
    )

    doc.add_heading("3.2 归并排序", level=2)
    doc.add_paragraph(
        "归并排序同样采用分治策略，将数组递归拆分为单元素子数组，"
        "然后逐步合并。时间复杂度稳定 O(n log n)，空间 O(n)。"
    )

    doc.add_heading("四、实验结果", level=1)
    table = doc.add_table(rows=6, cols=5)
    table.style = "Table Grid"
    headers = ["数据规模", "冒泡排序(ms)", "快速排序(ms)", "归并排序(ms)", "堆排序(ms)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    results = [
        ["1,000", "12.3", "0.8", "1.1", "1.3"],
        ["5,000", "298.5", "4.2", "5.8", "6.1"],
        ["10,000", "1205.7", "8.9", "12.3", "13.5"],
        ["50,000", "超时", "48.2", "65.1", "72.8"],
        ["100,000", "超时", "102.5", "138.7", "155.2"],
    ]
    for row_idx, row_data in enumerate(results, 1):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = val

    doc.add_heading("五、结论", level=1)
    doc.add_paragraph(
        "实验结果表明：\n"
        "1. 冒泡排序在大规模数据下性能极差，不适合实际应用\n"
        "2. 快速排序平均性能最优，但最坏情况需要优化（三数取中法）\n"
        "3. 归并排序性能稳定，适合外部排序场景\n"
        "4. 堆排序空间效率最高（原地排序），适合内存受限场景"
    )

    doc.save(str(OUTPUT_DIR / "sample_algorithm.docx"))
    print("✓ sample_algorithm.docx")


def generate_docx_api_test() -> None:
    """生成 API 测试报告 Word 文档."""
    from docx import Document

    doc = Document()
    doc.add_heading("RESTful API 接口测试报告", level=0)
    doc.add_paragraph("项目：学生成绩管理系统  版本：v2.1.0")
    doc.add_paragraph("测试人员：赵强  测试日期：2024-12-10")
    doc.add_paragraph("")

    doc.add_heading("一、测试概述", level=1)
    doc.add_paragraph(
        "本次测试覆盖学生成绩管理系统的全部 REST API 接口，"
        "共 28 个接口，使用 Postman + Newman 自动化执行。"
    )

    doc.add_heading("二、测试环境", level=1)
    doc.add_paragraph("- 服务端：Spring Boot 3.2 + MySQL 8.0")
    doc.add_paragraph("- 测试工具：Postman v10.21 + Newman CLI")
    doc.add_paragraph("- 测试数据：使用 Faker 生成 500 条学生记录")

    doc.add_heading("三、测试用例", level=1)
    doc.add_heading("3.1 用户认证接口", level=2)

    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["用例ID", "接口", "预期", "结果"]):
        table.rows[0].cells[i].text = h
    cases = [
        ["TC-001", "POST /api/login", "200 + token", "通过"],
        ["TC-002", "POST /api/login (错误密码)", "401", "通过"],
        ["TC-003", "GET /api/profile (无token)", "401", "通过"],
        ["TC-004", "POST /api/register", "201", "通过"],
    ]
    for row_idx, row_data in enumerate(cases, 1):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = val

    doc.add_heading("3.2 成绩管理接口", level=2)
    table2 = doc.add_table(rows=6, cols=4)
    table2.style = "Table Grid"
    for i, h in enumerate(["用例ID", "接口", "预期", "结果"]):
        table2.rows[0].cells[i].text = h
    cases2 = [
        ["TC-010", "GET /api/scores?page=1", "200 + 分页数据", "通过"],
        ["TC-011", "POST /api/scores", "201", "通过"],
        ["TC-012", "PUT /api/scores/1", "200", "通过"],
        ["TC-013", "DELETE /api/scores/1", "204", "通过"],
        ["TC-014", "GET /api/scores/stats", "200 + 统计", "通过"],
    ]
    for row_idx, row_data in enumerate(cases2, 1):
        for col_idx, val in enumerate(row_data):
            table2.rows[row_idx].cells[col_idx].text = val

    doc.add_heading("四、测试结论", level=1)
    doc.add_paragraph("通过率：28/28 = 100%")
    doc.add_paragraph("平均响应时间：45ms")
    doc.add_paragraph("所有接口均符合 RESTful 规范，错误处理完善。")

    doc.save(str(OUTPUT_DIR / "sample_api_test.docx"))
    print("✓ sample_api_test.docx")


def generate_pdf_database() -> None:
    """生成数据库设计文档 PDF."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    pdf_path = str(OUTPUT_DIR / "sample_database_design.pdf")
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    styles = getSampleStyleSheet()

    # 注册中文字体（如果可用）
    chinese_style = ParagraphStyle(
        "Chinese", parent=styles["Normal"], fontName="Helvetica", fontSize=11
    )
    title_style = ParagraphStyle(
        "ChTitle", parent=styles["Heading1"], fontName="Helvetica", fontSize=18
    )
    h2_style = ParagraphStyle(
        "ChH2", parent=styles["Heading2"], fontName="Helvetica", fontSize=14
    )

    elements = []
    elements.append(Paragraph("Database Design Document", title_style))
    elements.append(Paragraph("Online Shopping System - E-Commerce Platform", chinese_style))
    elements.append(Spacer(1, 1 * cm))

    elements.append(Paragraph("1. Overview", h2_style))
    elements.append(Paragraph(
        "This document describes the database schema for an e-commerce platform. "
        "The system uses MySQL 8.0 with InnoDB engine for ACID compliance. "
        "The schema supports user management, product catalog, order processing, "
        "and payment tracking.",
        chinese_style,
    ))
    elements.append(Spacer(1, 0.5 * cm))

    elements.append(Paragraph("2. Entity Relationship", h2_style))
    elements.append(Paragraph(
        "Core entities: Users, Products, Categories, Orders, OrderItems, "
        "Payments, Reviews, Addresses. Relationships follow 3NF normalization.",
        chinese_style,
    ))
    elements.append(Spacer(1, 0.5 * cm))

    # Users table
    elements.append(Paragraph("3. Table Definitions", h2_style))
    elements.append(Paragraph("3.1 Users Table", chinese_style))
    user_data = [
        ["Column", "Type", "Constraints", "Description"],
        ["id", "BIGINT", "PK, AUTO_INCREMENT", "User ID"],
        ["username", "VARCHAR(50)", "UNIQUE, NOT NULL", "Login name"],
        ["email", "VARCHAR(100)", "UNIQUE, NOT NULL", "Email address"],
        ["password_hash", "VARCHAR(255)", "NOT NULL", "Bcrypt hash"],
        ["phone", "VARCHAR(20)", "NULL", "Phone number"],
        ["status", "ENUM", "DEFAULT 'active'", "active/disabled/banned"],
        ["created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Registration time"],
    ]
    t = Table(user_data, colWidths=[3 * cm, 3 * cm, 4 * cm, 4 * cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 0.5 * cm))

    # Products table
    elements.append(Paragraph("3.2 Products Table", chinese_style))
    prod_data = [
        ["Column", "Type", "Constraints", "Description"],
        ["id", "BIGINT", "PK, AUTO_INCREMENT", "Product ID"],
        ["name", "VARCHAR(200)", "NOT NULL", "Product name"],
        ["description", "TEXT", "NULL", "Full description"],
        ["price", "DECIMAL(10,2)", "NOT NULL, CHECK > 0", "Unit price"],
        ["stock", "INT", "DEFAULT 0, CHECK >= 0", "Available stock"],
        ["category_id", "BIGINT", "FK -> categories.id", "Category"],
        ["is_active", "BOOLEAN", "DEFAULT TRUE", "Listing status"],
    ]
    t2 = Table(prod_data, colWidths=[3 * cm, 3 * cm, 4 * cm, 4 * cm])
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
    ]))
    elements.append(t2)
    elements.append(Spacer(1, 0.5 * cm))

    # Indexes section
    elements.append(Paragraph("4. Index Strategy", h2_style))
    elements.append(Paragraph(
        "- B-Tree index on users.email for login queries\n"
        "- Composite index on orders(user_id, status) for order listing\n"
        "- Full-text index on products.name + products.description for search\n"
        "- Covering index on order_items(order_id, product_id, quantity, price)",
        chinese_style,
    ))
    elements.append(Spacer(1, 0.5 * cm))

    elements.append(Paragraph("5. SQL Scripts", h2_style))
    elements.append(Paragraph(
        "CREATE TABLE users (\n"
        "  id BIGINT AUTO_INCREMENT PRIMARY KEY,\n"
        "  username VARCHAR(50) NOT NULL UNIQUE,\n"
        "  email VARCHAR(100) NOT NULL UNIQUE,\n"
        "  password_hash VARCHAR(255) NOT NULL,\n"
        "  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP\n"
        ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;",
        chinese_style,
    ))

    doc.build(elements)
    print("✓ sample_database_design.pdf")


def generate_excel_data_analysis() -> None:
    """生成数据分析实验 Excel 表格."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    wb = Workbook()

    # Sheet 1: 实验数据
    ws1 = wb.active
    ws1.title = "Sales Data"
    headers = ["Month", "Product", "Region", "Units Sold", "Revenue", "Cost", "Profit"]
    for col, h in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="4472C4", fill_type="solid")
        cell.font = Font(bold=True, color="FFFFFF")

    import random
    random.seed(42)
    products = ["Laptop", "Phone", "Tablet", "Monitor", "Keyboard"]
    regions = ["North", "South", "East", "West"]
    months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
              "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]

    row = 2
    for month in months:
        for product in products:
            for region in regions:
                units = random.randint(10, 200)
                price = {"Laptop": 899, "Phone": 699, "Tablet": 499,
                         "Monitor": 349, "Keyboard": 79}[product]
                revenue = units * price
                cost = int(revenue * random.uniform(0.5, 0.7))
                profit = revenue - cost
                ws1.cell(row=row, column=1, value=month)
                ws1.cell(row=row, column=2, value=product)
                ws1.cell(row=row, column=3, value=region)
                ws1.cell(row=row, column=4, value=units)
                ws1.cell(row=row, column=5, value=revenue)
                ws1.cell(row=row, column=6, value=cost)
                ws1.cell(row=row, column=7, value=profit)
                row += 1

    # Sheet 2: 统计汇总
    ws2 = wb.create_sheet("Summary Statistics")
    ws2.cell(row=1, column=1, value="Metric").font = Font(bold=True)
    ws2.cell(row=1, column=2, value="Value").font = Font(bold=True)
    stats = [
        ("Total Records", row - 2),
        ("Total Revenue", "=SUM('Sales Data'!E2:E241)"),
        ("Total Profit", "=SUM('Sales Data'!G2:G241)"),
        ("Avg Units/Record", "=AVERAGE('Sales Data'!D2:D241)"),
        ("Max Single Sale", "=MAX('Sales Data'!E2:E241)"),
        ("Min Single Sale", "=MIN('Sales Data'!E2:E241)"),
        ("Profit Margin %", "=SUM('Sales Data'!G2:G241)/SUM('Sales Data'!E2:E241)*100"),
    ]
    for i, (metric, value) in enumerate(stats, 2):
        ws2.cell(row=i, column=1, value=metric)
        ws2.cell(row=i, column=2, value=value)

    # Sheet 3: 分析结论
    ws3 = wb.create_sheet("Analysis Notes")
    notes = [
        "Data Analysis Experiment - Conclusions",
        "",
        "1. Laptop category generates highest revenue (45% of total)",
        "2. North region consistently outperforms other regions",
        "3. Q4 (Oct-Dec) shows 30% increase vs Q1 due to holiday season",
        "4. Keyboard has highest profit margin (48%) despite lowest revenue",
        "5. Phone sales show declining trend - recommend marketing review",
        "",
        "Methodology:",
        "- Data collected from company ERP system (Jan-Dec 2024)",
        "- Statistical analysis using pandas + matplotlib",
        "- Visualization: bar charts, line plots, heatmaps",
        "- Hypothesis testing: t-test for regional differences (p<0.05)",
    ]
    for i, note in enumerate(notes, 1):
        ws3.cell(row=i, column=1, value=note)

    wb.save(str(OUTPUT_DIR / "sample_data_analysis.xlsx"))
    print("✓ sample_data_analysis.xlsx")


def generate_excel_ml_notebook() -> None:
    """生成机器学习实验数据 Excel."""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Model Comparison"

    headers = ["Model", "Accuracy", "Precision", "Recall", "F1-Score", "Train Time(s)", "Params"]
    for col, h in enumerate(headers, 1):
        ws1.cell(row=1, column=col, value=h).font = Font(bold=True)

    models = [
        ["Logistic Regression", 0.847, 0.832, 0.861, 0.846, 2.3, 785],
        ["Random Forest", 0.912, 0.905, 0.918, 0.911, 15.7, 50000],
        ["SVM (RBF)", 0.889, 0.876, 0.901, 0.888, 45.2, 12000],
        ["XGBoost", 0.923, 0.917, 0.929, 0.923, 8.9, 35000],
        ["Neural Network (MLP)", 0.908, 0.895, 0.921, 0.908, 120.5, 150000],
        ["KNN (k=5)", 0.834, 0.821, 0.847, 0.834, 0.1, 0],
        ["Naive Bayes", 0.791, 0.778, 0.804, 0.791, 0.5, 200],
    ]
    for row_idx, row_data in enumerate(models, 2):
        for col_idx, val in enumerate(row_data):
            ws1.cell(row=row_idx, column=col_idx + 1, value=val)

    # Sheet 2: Hyperparameter tuning
    ws2 = wb.create_sheet("XGBoost Tuning")
    headers2 = ["n_estimators", "max_depth", "learning_rate", "CV Score", "Std"]
    for col, h in enumerate(headers2, 1):
        ws2.cell(row=1, column=col, value=h).font = Font(bold=True)

    import random
    random.seed(123)
    row = 2
    for n_est in [50, 100, 200, 300, 500]:
        for depth in [3, 5, 7, 9]:
            for lr in [0.01, 0.05, 0.1, 0.2]:
                score = 0.85 + random.uniform(0, 0.08)
                std = random.uniform(0.005, 0.02)
                ws2.cell(row=row, column=1, value=n_est)
                ws2.cell(row=row, column=2, value=depth)
                ws2.cell(row=row, column=3, value=lr)
                ws2.cell(row=row, column=4, value=round(score, 4))
                ws2.cell(row=row, column=5, value=round(std, 4))
                row += 1

    # Sheet 3: Feature importance
    ws3 = wb.create_sheet("Feature Importance")
    features = [
        ("age", 0.156), ("income", 0.234), ("education_years", 0.189),
        ("work_experience", 0.145), ("credit_score", 0.312),
        ("num_dependents", 0.067), ("loan_amount", 0.278),
        ("employment_type", 0.098), ("property_area", 0.045),
        ("marital_status", 0.032),
    ]
    ws3.cell(row=1, column=1, value="Feature").font = Font(bold=True)
    ws3.cell(row=1, column=2, value="Importance").font = Font(bold=True)
    for i, (feat, imp) in enumerate(features, 2):
        ws3.cell(row=i, column=1, value=feat)
        ws3.cell(row=i, column=2, value=imp)

    wb.save(str(OUTPUT_DIR / "sample_ml_notebook.xlsx"))
    print("✓ sample_ml_notebook.xlsx")


def generate_zip_flask_project() -> None:
    """生成 Flask Web 项目源代码压缩包."""
    zip_path = str(OUTPUT_DIR / "sample_flask_project.zip")

    files = {
        "book_manager/app.py": '''"""Flask Application Entry Point."""
from flask import Flask
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager

db = SQLAlchemy()
login_manager = LoginManager()


def create_app(config_name="development"):
    app = Flask(__name__)
    app.config.from_object(f"config.{config_name}")

    db.init_app(app)
    login_manager.init_app(app)
    login_manager.login_view = "auth.login"

    from .routes.auth import auth_bp
    from .routes.books import books_bp
    from .routes.borrow import borrow_bp

    app.register_blueprint(auth_bp)
    app.register_blueprint(books_bp)
    app.register_blueprint(borrow_bp)

    return app
''',
        "book_manager/models.py": '''"""Database Models."""
from datetime import datetime, timedelta
from flask_login import UserMixin
from werkzeug.security import generate_password_hash, check_password_hash
from . import db


class User(UserMixin, db.Model):
    __tablename__ = "users"

    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.String(128), nullable=False)
    role = db.Column(db.String(20), default="reader")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    borrows = db.relationship("BorrowRecord", backref="user", lazy="dynamic")

    def set_password(self, password: str) -> None:
        self.password_hash = generate_password_hash(password)

    def check_password(self, password: str) -> bool:
        return check_password_hash(self.password_hash, password)


class Book(db.Model):
    __tablename__ = "books"

    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    author = db.Column(db.String(100), nullable=False)
    isbn = db.Column(db.String(13), unique=True)
    category_id = db.Column(db.Integer, db.ForeignKey("categories.id"))
    total_copies = db.Column(db.Integer, default=1)
    available_copies = db.Column(db.Integer, default=1)
    description = db.Column(db.Text)
    published_year = db.Column(db.Integer)

    category = db.relationship("Category", backref="books")


class Category(db.Model):
    __tablename__ = "categories"

    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50), unique=True, nullable=False)
    description = db.Column(db.String(200))


class BorrowRecord(db.Model):
    __tablename__ = "borrow_records"

    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    book_id = db.Column(db.Integer, db.ForeignKey("books.id"), nullable=False)
    borrow_date = db.Column(db.DateTime, default=datetime.utcnow)
    due_date = db.Column(db.DateTime)
    return_date = db.Column(db.DateTime, nullable=True)
    status = db.Column(db.String(20), default="borrowed")

    book = db.relationship("Book", backref="borrow_records")

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        if not self.due_date:
            self.due_date = datetime.utcnow() + timedelta(days=30)

    @property
    def is_overdue(self) -> bool:
        if self.status == "returned":
            return False
        return datetime.utcnow() > self.due_date
''',
    }

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path, content in files.items():
            zf.writestr(path, content)

    # Add more files to the zip
    more_files = {
        "book_manager/routes/__init__.py": "",
        "book_manager/routes/auth.py": '''"""Authentication Routes."""
from flask import Blueprint, request, jsonify
from flask_login import login_user, logout_user, login_required, current_user
from ..models import User, db

auth_bp = Blueprint("auth", __name__, url_prefix="/api/auth")


@auth_bp.route("/register", methods=["POST"])
def register():
    data = request.get_json()
    if not data:
        return jsonify({"error": "No data provided"}), 400

    username = data.get("username", "").strip()
    email = data.get("email", "").strip()
    password = data.get("password", "")

    if not username or not email or not password:
        return jsonify({"error": "All fields required"}), 400

    if len(password) < 8:
        return jsonify({"error": "Password must be at least 8 chars"}), 400

    if User.query.filter_by(username=username).first():
        return jsonify({"error": "Username already exists"}), 409

    user = User(username=username, email=email)
    user.set_password(password)
    db.session.add(user)
    db.session.commit()

    return jsonify({"id": user.id, "username": user.username}), 201


@auth_bp.route("/login", methods=["POST"])
def login():
    data = request.get_json()
    username = data.get("username", "")
    password = data.get("password", "")

    user = User.query.filter_by(username=username).first()
    if user is None or not user.check_password(password):
        return jsonify({"error": "Invalid credentials"}), 401

    login_user(user)
    return jsonify({"message": "Login successful", "user_id": user.id})


@auth_bp.route("/logout", methods=["POST"])
@login_required
def logout():
    logout_user()
    return jsonify({"message": "Logged out"})
''',
        "book_manager/routes/books.py": '''"""Book Management Routes."""
from flask import Blueprint, request, jsonify
from flask_login import login_required
from ..models import Book, Category, db

books_bp = Blueprint("books", __name__, url_prefix="/api/books")


@books_bp.route("/", methods=["GET"])
def list_books():
    page = request.args.get("page", 1, type=int)
    per_page = request.args.get("per_page", 20, type=int)
    query = request.args.get("q", "")

    books_query = Book.query
    if query:
        books_query = books_query.filter(
            Book.title.ilike(f"%{query}%") | Book.author.ilike(f"%{query}%")
        )

    pagination = books_query.paginate(page=page, per_page=per_page)
    return jsonify({
        "items": [
            {
                "id": b.id,
                "title": b.title,
                "author": b.author,
                "isbn": b.isbn,
                "available": b.available_copies,
                "category": b.category.name if b.category else None,
            }
            for b in pagination.items
        ],
        "total": pagination.total,
        "pages": pagination.pages,
        "current_page": page,
    })


@books_bp.route("/", methods=["POST"])
@login_required
def create_book():
    data = request.get_json()
    book = Book(
        title=data["title"],
        author=data["author"],
        isbn=data.get("isbn"),
        total_copies=data.get("total_copies", 1),
        available_copies=data.get("total_copies", 1),
        description=data.get("description"),
        category_id=data.get("category_id"),
    )
    db.session.add(book)
    db.session.commit()
    return jsonify({"id": book.id, "title": book.title}), 201
''',
        "book_manager/routes/borrow.py": '''"""Borrow Management Routes."""
from datetime import datetime
from flask import Blueprint, request, jsonify
from flask_login import login_required, current_user
from ..models import Book, BorrowRecord, db

borrow_bp = Blueprint("borrow", __name__, url_prefix="/api/borrow")


@borrow_bp.route("/", methods=["POST"])
@login_required
def borrow_book():
    data = request.get_json()
    book_id = data.get("book_id")

    book = Book.query.get_or_404(book_id)
    if book.available_copies <= 0:
        return jsonify({"error": "No copies available"}), 400

    # Check if user already has this book
    existing = BorrowRecord.query.filter_by(
        user_id=current_user.id, book_id=book_id, status="borrowed"
    ).first()
    if existing:
        return jsonify({"error": "Already borrowed this book"}), 400

    record = BorrowRecord(user_id=current_user.id, book_id=book_id)
    book.available_copies -= 1
    db.session.add(record)
    db.session.commit()

    return jsonify({
        "id": record.id,
        "book": book.title,
        "due_date": record.due_date.isoformat(),
    }), 201


@borrow_bp.route("/<int:record_id>/return", methods=["POST"])
@login_required
def return_book(record_id):
    record = BorrowRecord.query.get_or_404(record_id)
    if record.user_id != current_user.id:
        return jsonify({"error": "Not your borrow record"}), 403
    if record.status == "returned":
        return jsonify({"error": "Already returned"}), 400

    record.status = "returned"
    record.return_date = datetime.utcnow()
    record.book.available_copies += 1
    db.session.commit()

    return jsonify({"message": "Book returned successfully"})
''',
    }

    with zipfile.ZipFile(zip_path, "a", zipfile.ZIP_DEFLATED) as zf:
        for path, content in more_files.items():
            zf.writestr(path, content)

    # Add config and tests
    config_files = {
        "config.py": '''"""Application Configuration."""
import os


class development:
    SECRET_KEY = os.environ.get("SECRET_KEY", "dev-secret-key")
    SQLALCHEMY_DATABASE_URI = "mysql+pymysql://root:password@localhost/bookdb"
    SQLALCHEMY_TRACK_MODIFICATIONS = False


class testing:
    TESTING = True
    SQLALCHEMY_DATABASE_URI = "sqlite:///:memory:"
    SECRET_KEY = "test-secret"


class production:
    SECRET_KEY = os.environ["SECRET_KEY"]
    SQLALCHEMY_DATABASE_URI = os.environ["DATABASE_URL"]
    SQLALCHEMY_TRACK_MODIFICATIONS = False
''',
        "requirements.txt": "flask==2.3.2\nflask-sqlalchemy==3.1.1\nflask-login==0.6.3\n"
                           "pymysql==1.1.0\nwerkzeug==2.3.7\npytest==7.4.3\n",
        "tests/test_auth.py": '''"""Authentication Tests."""
import pytest
from book_manager import create_app, db
from book_manager.models import User


@pytest.fixture
def app():
    app = create_app("testing")
    with app.app_context():
        db.create_all()
        yield app
        db.drop_all()


@pytest.fixture
def client(app):
    return app.test_client()


def test_register(client):
    resp = client.post("/api/auth/register", json={
        "username": "testuser",
        "email": "test@example.com",
        "password": "securepass123",
    })
    assert resp.status_code == 201
    data = resp.get_json()
    assert data["username"] == "testuser"


def test_register_duplicate(client):
    client.post("/api/auth/register", json={
        "username": "testuser",
        "email": "test@example.com",
        "password": "securepass123",
    })
    resp = client.post("/api/auth/register", json={
        "username": "testuser",
        "email": "test2@example.com",
        "password": "securepass123",
    })
    assert resp.status_code == 409


def test_login_success(client):
    client.post("/api/auth/register", json={
        "username": "testuser",
        "email": "test@example.com",
        "password": "securepass123",
    })
    resp = client.post("/api/auth/login", json={
        "username": "testuser",
        "password": "securepass123",
    })
    assert resp.status_code == 200


def test_login_wrong_password(client):
    client.post("/api/auth/register", json={
        "username": "testuser",
        "email": "test@example.com",
        "password": "securepass123",
    })
    resp = client.post("/api/auth/login", json={
        "username": "testuser",
        "password": "wrongpassword",
    })
    assert resp.status_code == 401
''',
        "README.md": "# Book Manager\n\nA Flask-based library management system.\n\n"
                    "## Setup\n```bash\npip install -r requirements.txt\n"
                    "flask db upgrade\nflask run\n```\n\n"
                    "## Features\n- User authentication\n- Book CRUD\n"
                    "- Borrow/Return management\n- Search functionality\n",
    }

    with zipfile.ZipFile(zip_path, "a", zipfile.ZIP_DEFLATED) as zf:
        for path, content in config_files.items():
            zf.writestr(path, content)

    print("✓ sample_flask_project.zip")


def generate_zip_vue_project() -> None:
    """生成 Vue.js 前端项目源代码压缩包."""
    zip_path = str(OUTPUT_DIR / "sample_vue_project.zip")

    files = {
        "todo-app/package.json": '''{
  "name": "todo-app",
  "version": "1.0.0",
  "scripts": {
    "dev": "vite",
    "build": "vite build",
    "test": "vitest"
  },
  "dependencies": {
    "vue": "^3.4.0",
    "vue-router": "^4.2.0",
    "pinia": "^2.1.0",
    "axios": "^1.6.0"
  },
  "devDependencies": {
    "vite": "^5.0.0",
    "@vitejs/plugin-vue": "^5.0.0",
    "vitest": "^1.0.0",
    "tailwindcss": "^3.4.0"
  }
}
''',
        "todo-app/src/main.ts": '''import { createApp } from "vue"
import { createPinia } from "pinia"
import App from "./App.vue"
import router from "./router"
import "./style.css"

const app = createApp(App)
app.use(createPinia())
app.use(router)
app.mount("#app")
''',
        "todo-app/src/App.vue": '''<template>
  <div class="min-h-screen bg-gray-100">
    <nav class="bg-white shadow-sm">
      <div class="max-w-4xl mx-auto px-4 py-3">
        <h1 class="text-xl font-bold text-gray-800">Todo App</h1>
      </div>
    </nav>
    <main class="max-w-4xl mx-auto px-4 py-8">
      <router-view />
    </main>
  </div>
</template>
''',
        "todo-app/src/router/index.ts": '''import { createRouter, createWebHistory } from "vue-router"

const router = createRouter({
  history: createWebHistory(),
  routes: [
    { path: "/", component: () => import("../views/HomeView.vue") },
    { path: "/about", component: () => import("../views/AboutView.vue") },
  ],
})

export default router
''',
    }

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path, content in files.items():
            zf.writestr(path, content)

    more_vue_files = {
        "todo-app/src/stores/todo.ts": '''import { defineStore } from "pinia"
import { ref, computed } from "vue"

export interface Todo {
  id: number
  title: string
  completed: boolean
  createdAt: Date
  priority: "low" | "medium" | "high"
}

export const useTodoStore = defineStore("todo", () => {
  const todos = ref<Todo[]>([])
  let nextId = 1

  const completedCount = computed(() =>
    todos.value.filter((t) => t.completed).length
  )

  const pendingCount = computed(() =>
    todos.value.filter((t) => !t.completed).length
  )

  function addTodo(title: string, priority: Todo["priority"] = "medium") {
    todos.value.push({
      id: nextId++,
      title,
      completed: false,
      createdAt: new Date(),
      priority,
    })
  }

  function toggleTodo(id: number) {
    const todo = todos.value.find((t) => t.id === id)
    if (todo) {
      todo.completed = !todo.completed
    }
  }

  function removeTodo(id: number) {
    todos.value = todos.value.filter((t) => t.id !== id)
  }

  function clearCompleted() {
    todos.value = todos.value.filter((t) => !t.completed)
  }

  return { todos, completedCount, pendingCount, addTodo, toggleTodo, removeTodo, clearCompleted }
})
''',
        "todo-app/src/views/HomeView.vue": '''<script setup lang="ts">
import { ref } from "vue"
import { useTodoStore } from "../stores/todo"
import TodoItem from "../components/TodoItem.vue"

const store = useTodoStore()
const newTitle = ref("")
const newPriority = ref<"low" | "medium" | "high">("medium")

function handleAdd() {
  if (newTitle.value.trim()) {
    store.addTodo(newTitle.value.trim(), newPriority.value)
    newTitle.value = ""
  }
}
</script>

<template>
  <div class="space-y-6">
    <div class="bg-white rounded-lg shadow p-6">
      <h2 class="text-lg font-semibold mb-4">Add New Todo</h2>
      <form @submit.prevent="handleAdd" class="flex gap-3">
        <input
          v-model="newTitle"
          type="text"
          placeholder="What needs to be done?"
          class="flex-1 border rounded px-3 py-2"
        />
        <select v-model="newPriority" class="border rounded px-3 py-2">
          <option value="low">Low</option>
          <option value="medium">Medium</option>
          <option value="high">High</option>
        </select>
        <button type="submit" class="bg-blue-500 text-white px-4 py-2 rounded">
          Add
        </button>
      </form>
    </div>

    <div class="bg-white rounded-lg shadow p-6">
      <div class="flex justify-between items-center mb-4">
        <h2 class="text-lg font-semibold">
          Todos ({{ store.pendingCount }} pending)
        </h2>
        <button
          v-if="store.completedCount > 0"
          @click="store.clearCompleted()"
          class="text-sm text-red-500"
        >
          Clear completed ({{ store.completedCount }})
        </button>
      </div>
      <div class="space-y-2">
        <TodoItem
          v-for="todo in store.todos"
          :key="todo.id"
          :todo="todo"
          @toggle="store.toggleTodo(todo.id)"
          @remove="store.removeTodo(todo.id)"
        />
        <p v-if="store.todos.length === 0" class="text-gray-400 text-center py-8">
          No todos yet. Add one above!
        </p>
      </div>
    </div>
  </div>
</template>
''',
        "todo-app/src/components/TodoItem.vue": '''<script setup lang="ts">
import type { Todo } from "../stores/todo"

defineProps<{ todo: Todo }>()
defineEmits<{ toggle: []; remove: [] }>()

const priorityColors = {
  low: "bg-green-100 text-green-800",
  medium: "bg-yellow-100 text-yellow-800",
  high: "bg-red-100 text-red-800",
}
</script>

<template>
  <div class="flex items-center gap-3 p-3 border rounded hover:bg-gray-50">
    <input
      type="checkbox"
      :checked="todo.completed"
      @change="$emit(\'toggle\')"
      class="w-5 h-5"
    />
    <span :class="{ \'line-through text-gray-400\': todo.completed }" class="flex-1">
      {{ todo.title }}
    </span>
    <span :class="priorityColors[todo.priority]" class="text-xs px-2 py-1 rounded">
      {{ todo.priority }}
    </span>
    <button @click="$emit(\'remove\')" class="text-red-400 hover:text-red-600">
      &times;
    </button>
  </div>
</template>
''',
    }

    with zipfile.ZipFile(zip_path, "a", zipfile.ZIP_DEFLATED) as zf:
        for path, content in more_vue_files.items():
            zf.writestr(path, content)

    print("✓ sample_vue_project.zip")


def generate_png_screenshot() -> None:
    """生成一张含中文文字的程序运行截图（使用 Pillow 绘制）."""
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        print("⚠ Pillow not installed, skipping PNG generation")
        return

    # 创建一个模拟的终端/IDE 截图
    width, height = 800, 600
    img = Image.new("RGB", (width, height), color=(30, 30, 30))
    draw = ImageDraw.Draw(img)

    # 尝试使用系统字体
    try:
        font = ImageFont.truetype("consola.ttf", 14)
        font_title = ImageFont.truetype("consola.ttf", 16)
    except (OSError, IOError):
        font = ImageFont.load_default()
        font_title = font

    # 绘制标题栏
    draw.rectangle([(0, 0), (width, 30)], fill=(50, 50, 50))
    draw.text((10, 7), "Terminal - Python 3.11", fill=(200, 200, 200), font=font_title)

    # 绘制终端内容
    lines = [
        "$ python manage.py test",
        "Running tests...",
        "test_auth.test_register ... OK",
        "test_auth.test_login ... OK",
        "test_auth.test_logout ... OK",
        "test_books.test_list_books ... OK",
        "test_books.test_create_book ... OK",
        "test_books.test_search ... OK",
        "test_borrow.test_borrow_book ... OK",
        "test_borrow.test_return_book ... OK",
        "test_borrow.test_overdue_check ... OK",
        "",
        "----------------------------------------------",
        "Ran 9 tests in 2.341s",
        "",
        "OK",
        "",
        "$ python manage.py runserver",
        " * Running on http://127.0.0.1:5000",
        " * Debug mode: on",
        "",
        "Database: MySQL 5.7 connected",
        "Redis: connected (localhost:6379)",
        "Server ready. Accepting connections...",
    ]

    y = 40
    for line in lines:
        color = (0, 255, 0) if "OK" in line and "..." in line else (200, 200, 200)
        if line.startswith("$"):
            color = (100, 200, 255)
        if "ERROR" in line or "FAIL" in line:
            color = (255, 80, 80)
        draw.text((15, y), line, fill=color, font=font)
        y += 22

    img.save(str(OUTPUT_DIR / "sample_screenshot.png"))
    print("✓ sample_screenshot.png")


def generate_pdf_network() -> None:
    """生成网络配置文档 PDF（纯文本 PDF，不依赖 reportlab）."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        from reportlab.lib.units import cm
    except ImportError:
        print("⚠ reportlab not installed, skipping network PDF")
        return

    pdf_path = str(OUTPUT_DIR / "sample_network_config.pdf")
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    styles = getSampleStyleSheet()

    elements = []
    elements.append(Paragraph("Network Configuration Lab Report", styles["Heading1"]))
    elements.append(Paragraph("Student: Wang Lei | ID: 2024003", styles["Normal"]))
    elements.append(Spacer(1, 0.5 * cm))

    elements.append(Paragraph("1. Lab Objectives", styles["Heading2"]))
    elements.append(Paragraph(
        "Configure a small enterprise network with VLANs, inter-VLAN routing, "
        "DHCP, and basic ACLs using Cisco Packet Tracer.",
        styles["Normal"],
    ))
    elements.append(Spacer(1, 0.3 * cm))

    elements.append(Paragraph("2. Network Topology", styles["Heading2"]))
    elements.append(Paragraph(
        "The network consists of: 1x Layer 3 Switch (Core), "
        "2x Layer 2 Switches (Access), 1x Router (Gateway), "
        "3x VLANs (Management=10, Staff=20, Guest=30), "
        "DHCP Server on VLAN 10.",
        styles["Normal"],
    ))
    elements.append(Spacer(1, 0.3 * cm))

    elements.append(Paragraph("3. Configuration Commands", styles["Heading2"]))
    elements.append(Paragraph("3.1 VLAN Configuration (Core Switch)", styles["Heading3"]))
    config_text = (
        "Switch(config)# vlan 10<br/>"
        "Switch(config-vlan)# name Management<br/>"
        "Switch(config)# vlan 20<br/>"
        "Switch(config-vlan)# name Staff<br/>"
        "Switch(config)# vlan 30<br/>"
        "Switch(config-vlan)# name Guest<br/>"
        "Switch(config)# interface vlan 10<br/>"
        "Switch(config-if)# ip address 192.168.10.1 255.255.255.0<br/>"
        "Switch(config-if)# no shutdown<br/>"
        "Switch(config)# interface vlan 20<br/>"
        "Switch(config-if)# ip address 192.168.20.1 255.255.255.0<br/>"
        "Switch(config-if)# no shutdown<br/>"
        "Switch(config)# interface vlan 30<br/>"
        "Switch(config-if)# ip address 192.168.30.1 255.255.255.0<br/>"
        "Switch(config-if)# no shutdown<br/>"
        "Switch(config)# ip routing"
    )
    elements.append(Paragraph(config_text, styles["Code"]))
    elements.append(Spacer(1, 0.3 * cm))

    elements.append(Paragraph("3.2 DHCP Configuration", styles["Heading3"]))
    dhcp_text = (
        "Router(config)# ip dhcp pool STAFF<br/>"
        "Router(dhcp-config)# network 192.168.20.0 255.255.255.0<br/>"
        "Router(dhcp-config)# default-router 192.168.20.1<br/>"
        "Router(dhcp-config)# dns-server 8.8.8.8<br/>"
        "Router(config)# ip dhcp pool GUEST<br/>"
        "Router(dhcp-config)# network 192.168.30.0 255.255.255.0<br/>"
        "Router(dhcp-config)# default-router 192.168.30.1<br/>"
        "Router(dhcp-config)# lease 0 4"
    )
    elements.append(Paragraph(dhcp_text, styles["Code"]))
    elements.append(Spacer(1, 0.3 * cm))

    elements.append(Paragraph("3.3 ACL Configuration", styles["Heading3"]))
    acl_text = (
        "Router(config)# access-list 100 deny ip 192.168.30.0 0.0.0.255 192.168.10.0 0.0.0.255<br/>"
        "Router(config)# access-list 100 permit ip any any<br/>"
        "Router(config)# interface vlan 30<br/>"
        "Router(config-if)# ip access-group 100 in"
    )
    elements.append(Paragraph(acl_text, styles["Code"]))
    elements.append(Spacer(1, 0.3 * cm))

    elements.append(Paragraph("4. Verification Results", styles["Heading2"]))
    elements.append(Paragraph(
        "- Ping from VLAN 20 to VLAN 10: SUCCESS (avg 2ms)<br/>"
        "- Ping from VLAN 30 to VLAN 10: BLOCKED (ACL working)<br/>"
        "- DHCP lease on VLAN 20: 192.168.20.101 assigned<br/>"
        "- Inter-VLAN routing: All permitted paths working<br/>"
        "- Guest isolation: Confirmed no access to management VLAN",
        styles["Normal"],
    ))
    elements.append(Spacer(1, 0.3 * cm))

    elements.append(Paragraph("5. Conclusion", styles["Heading2"]))
    elements.append(Paragraph(
        "Successfully configured enterprise network with VLAN segmentation, "
        "inter-VLAN routing via L3 switch, DHCP for dynamic addressing, "
        "and ACLs for guest network isolation. All verification tests passed.",
        styles["Normal"],
    ))

    doc.build(elements)
    print("✓ sample_network_config.pdf")


def main() -> None:
    """生成所有测试样本."""
    print(f"Output directory: {OUTPUT_DIR}")
    print("=" * 50)

    # Word 文档（3 个不同类型）
    generate_docx_report()
    generate_docx_algorithm()
    generate_docx_api_test()

    # PDF 文档（2 个）
    generate_pdf_database()
    generate_pdf_network()

    # Excel 表格（2 个）
    generate_excel_data_analysis()
    generate_excel_ml_notebook()

    # 源代码压缩包（2 个）
    generate_zip_flask_project()
    generate_zip_vue_project()

    # 图片截图（1 个）
    generate_png_screenshot()

    print("=" * 50)
    print(f"All samples generated in: {OUTPUT_DIR}")
    # 列出生成的文件
    for f in sorted(OUTPUT_DIR.iterdir()):
        size_kb = f.stat().st_size / 1024
        print(f"  {f.name:40s} {size_kb:8.1f} KB")


if __name__ == "__main__":
    main()
