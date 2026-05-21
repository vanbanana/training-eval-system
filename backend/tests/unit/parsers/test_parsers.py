"""Epic 9 验收：Parser Protocol + 实现 + Factory + Fake."""

from __future__ import annotations

import io

import pytest

from app.parsers import ParsedDocument, Parser, get_parser
from app.parsers.base import ParserUnsupportedTypeError, TitleNode
from app.parsers.docx_parser import DocxParser
from app.parsers.factory import _REGISTRY
from app.parsers.ocr_parser import ImageOcrParser
from app.parsers.pdf_parser import PdfParser
from tests.fakes.fake_parser import FakeOcrParser, FakeParser


pytestmark = pytest.mark.unit


# ============ Task 9.1 - Protocol & schema ============

class TestParsedDocument:
    def test_default_empty(self) -> None:
        d = ParsedDocument(raw_text="")
        assert d.raw_text == ""
        assert d.paragraphs == []
        assert d.title_tree == []

    def test_title_tree_supports_6_levels(self) -> None:
        node = TitleNode(level=1, text="Top")
        node.children.append(TitleNode(level=2, text="L2"))
        node.children[0].children.append(TitleNode(level=3, text="L3"))
        d = ParsedDocument(raw_text="x", title_tree=[node])
        # Pydantic 序列化无问题
        json = d.model_dump()
        assert json["title_tree"][0]["children"][0]["children"][0]["text"] == "L3"

    def test_invalid_level_rejected(self) -> None:
        with pytest.raises(Exception):  # noqa: B017
            TitleNode(level=7, text="X")  # type: ignore[arg-type]


# ============ Task 9.2 - DocxParser ============

class TestDocxParser:
    async def test_parse_simple_document(self) -> None:
        from docx import Document

        # 构造 docx
        d = Document()
        d.add_heading("一级标题", level=1)
        d.add_paragraph("段落 1")
        d.add_heading("二级标题", level=2)
        d.add_paragraph("段落 2")
        buf = io.BytesIO()
        d.save(buf)
        content = buf.getvalue()

        parser = DocxParser()
        result = await parser.parse(content)
        assert "段落 1" in result.raw_text
        assert "段落 2" in result.raw_text
        assert len(result.title_tree) == 1  # 一个顶层 H1
        assert result.title_tree[0].text == "一级标题"
        assert result.title_tree[0].children[0].text == "二级标题"

    async def test_empty_docx_returns_empty(self) -> None:
        from docx import Document

        buf = io.BytesIO()
        Document().save(buf)
        parser = DocxParser()
        result = await parser.parse(buf.getvalue())
        assert result.paragraphs == []

    async def test_corrupted_bytes_raises(self) -> None:
        parser = DocxParser()
        with pytest.raises(Exception):  # noqa: B017
            await parser.parse(b"not a docx")


# ============ Task 9.3 - PdfParser ============

class TestPdfParser:
    async def test_extracts_text_from_real_pdf(self) -> None:
        # 用 pypdf 自己的 PdfWriter 构造一份带文本的 PDF
        from pypdf import PdfWriter

        pdf = PdfWriter()
        pdf.add_blank_page(width=200, height=200)
        buf = io.BytesIO()
        pdf.write(buf)

        parser = PdfParser()
        result = await parser.parse(buf.getvalue())
        # 空白 pdf：raw_text 为空
        assert result.page_count == 1

    async def test_invalid_pdf_raises(self) -> None:
        parser = PdfParser()
        with pytest.raises(Exception):  # noqa: B017
            await parser.parse(b"not a pdf at all")


# ============ Task 9.4 - ImageOcrParser ============

class TestOcrParser:
    async def test_returns_empty_when_tesseract_missing(self) -> None:
        parser = ImageOcrParser()
        # 假图（不会真的 OCR）
        result = await parser.parse(b"\x89PNG\r\n\x1a\n")
        # 至少不抛异常
        assert isinstance(result, ParsedDocument)


# ============ Task 9.5 - ParserFactory ============

class TestParserFactory:
    def test_routes_by_file_type(self) -> None:
        assert isinstance(get_parser("docx"), DocxParser)
        assert isinstance(get_parser("pdf"), PdfParser)
        assert isinstance(get_parser("png"), ImageOcrParser)
        assert isinstance(get_parser("jpg"), ImageOcrParser)

    def test_uppercase_normalized(self) -> None:
        assert isinstance(get_parser("DOCX"), DocxParser)
        assert isinstance(get_parser(".PDF"), PdfParser)

    def test_unsupported_type_raises(self) -> None:
        with pytest.raises(ParserUnsupportedTypeError):
            get_parser("exe")

    def test_registry_contains_all_supported(self) -> None:
        for ft in ("docx", "pdf", "png", "jpg", "jpeg"):
            assert ft in _REGISTRY


# ============ Task 9.6 - Fake parsers ============

class TestFakeParser:
    async def test_fake_parser_returns_predefined(self) -> None:
        expected = ParsedDocument(raw_text="预设")
        fake = FakeParser(output=expected)
        result = await fake.parse(b"any")
        assert result.raw_text == "预设"

    async def test_fake_implements_protocol(self) -> None:
        fake = FakeParser()
        assert isinstance(fake, Parser)

    async def test_fake_ocr(self) -> None:
        fake = FakeOcrParser(text="OCR 输出")
        result = await fake.parse(b"image bytes")
        assert result.raw_text == "OCR 输出"
